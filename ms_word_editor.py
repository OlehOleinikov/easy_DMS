"""
Формування документу MS Word
"""
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from person_builder import PersonProfile, DocumentId


class _DocEditorEmpty:
    """
    Клас, що містить екземпляр порожнього документу (python-docx) з типовими налаштуваннями та методами
    """

    def __init__(self):
        self.document = Document()
        self.sections = self.document.sections
        """Розмір полів"""
        for section in self.sections:
            section.page_height = Cm(29.7)
            section.page_width = Cm(21.0)
            section.top_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(1)
            section.bottom_margin = Cm(2)

        """Стилі тексту"""
        self.styles = self.document.styles
        # Стиль - заголовок "по центру":
        header_c = self.styles.add_style('central_header', WD_STYLE_TYPE.PARAGRAPH)
        header_c.font.name = 'Times New Roman'
        header_c.font.size = Pt(14)
        header_c.font.bold = True
        header_c.paragraph_format.space_before = Pt(0)
        header_c.paragraph_format.space_after = Pt(0)
        header_c.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Стиль - звичайний текст:
        text_base = self.styles.add_style('text_base', WD_STYLE_TYPE.PARAGRAPH)
        text_base.font.name = 'Times New Roman'
        text_base.font.size = Pt(14)
        text_base.paragraph_format.space_before = Pt(0)
        text_base.paragraph_format.space_after = Pt(0)
        text_base.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        text_base.paragraph_format.first_line_indent = Cm(1.25)
        text_base.paragraph_format.line_spacing = 1

        # Стиль - темно-червоний текст:
        text_red = self.styles.add_style('text_red', WD_STYLE_TYPE.CHARACTER)
        text_red.font.name = 'Times New Roman'
        text_red.font.size = Pt(14)
        text_red.font.color.rgb = RGBColor(127, 12, 7)

        # Дефолтний стиль - ненумерований список:
        list_style = self.styles['List Bullet']
        list_style.font.name = 'Times New Roman'
        list_style.font.size = Pt(14)
        list_style.paragraph_format.space_before = Pt(0)
        list_style.paragraph_format.space_after = Pt(0)
        list_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        list_style.paragraph_format.first_line_indent = Cm(1.25)
        list_style.paragraph_format.line_spacing = 1

        # Дефолтний стиль - ненумерований список другого порядку:
        list_style = self.styles['List Bullet 2']
        list_style.font.name = 'Times New Roman'
        list_style.font.size = Pt(14)
        list_style.paragraph_format.space_before = Pt(0)
        list_style.paragraph_format.space_after = Pt(0)
        list_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        list_style.paragraph_format.first_line_indent = Cm(-0.5)
        list_style.paragraph_format.left_indent = Cm(2.5)
        list_style.paragraph_format.line_spacing = 1

        tab_style = self.styles['Table Grid']
        tab_style.font.name = 'Times New Roman'
        tab_style.font.size = Pt(9)

    def save_docx(self, file_path):
        """Збереження документу python-docx у файл MS Word"""
        if type(file_path) == str:
            file_path = Path(file_path)
            file_path = file_path.with_suffix('.docx')
        try:
            self.document.save(file_path.absolute())
            return True
        except Exception:
            return False


class DocEditor(_DocEditorEmpty):
    """
    Клас формування документу з відомостями профайлу
    """
    def __init__(self):
        super().__init__()

    def add_card(self, card: PersonProfile):
        title = self.document.add_paragraph(f'', style='central_header')
        title.add_run(card.name_full, style='text_red').bold = True
        title.add_run(f', {card.dob} р.н.')

        if card.pob:
            p_pob = self.document.add_paragraph(f'Місце народження: {card.pob}.', style='text_base')

        if card.code_tax:
            p_tax = self.document.add_paragraph(f'РНОКПП: {card.code_tax}.', style='text_base')

        if card.code_uni:
            p_uni = self.document.add_paragraph(f'УНЗР: {card.code_uni}.', style='text_base')

        if card.address:
            p_adr = self.document.add_paragraph(f'Значиться за адресою: {card.address.strip()}.', style='text_base')

        if card.phone:
            p_adr = self.document.add_paragraph(f'Засіб зв\'язку ', style='text_base')
            p_adr.add_run(card.phone, style='text_red').bold = True
            p_adr.add_run(" (станом на момент перевірки перебуває на зв'язку).")

        if card.pass_internal:
            p_int_pass = self.document.add_paragraph(f'', style='text_base')
            p_int_pass.add_run(f'Паспорт громадянина України (внутрішній):').bold = True
            for passport in card.pass_internal:
                self._add_pass(passport)

        if card.pass_external:
            p_ext_pass = self.document.add_paragraph(f'', style='text_base')
            p_ext_pass.add_run(f'Паспорти громадянина України для виїзду за кордон:').bold = True
            for passport in card.pass_external:
                self._add_pass(passport)

        if card.image:
            p_image = self.document.add_paragraph(style='central_header')
            p_image.add_run().add_picture(card.image, width=Cm(4))
        self.document.add_page_break()

    def _add_pass(self, doc: DocumentId):
        p_pass = self.document.add_paragraph('', style='List Bullet')
        p_pass.add_run(doc.blank_number).bold = True
        if doc.date_created:
            p_pass.add_run(f' від {doc.date_created}')
        if doc.pass_office:
            p_pass.add_run(f' виданий: {doc.pass_office.strip()}')
        if doc.date_expired:
            p_pass.add_run(f', дійсний до: {doc.date_expired}')
        if doc.expired_status:
            p_pass.add_run(f' ({doc.expired_status.lower()})')
        p_pass.add_run('.')
